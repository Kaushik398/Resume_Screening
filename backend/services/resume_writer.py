import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_resume_docx(data: dict) -> bytes:
    doc = Document()
    
    # Adjust margins to fit on one or two pages nicely (0.75" margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Set default font styles to modern Calibri or Arial
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x11, 0x18, 0x27) # Dark charcoal text
    
    # Header: Full Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(data.get("full_name", "").upper())
    run_name.font.size = Pt(20)
    run_name.font.bold = True
    run_name.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89) # Professional deep blue accent
    p_name.paragraph_format.space_after = Pt(2)
    
    # Contact Info line
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = []
    if data.get("email"): contact_parts.append(data.get("email"))
    if data.get("phone"): contact_parts.append(data.get("phone"))
    if data.get("linkedin"): contact_parts.append(f"LinkedIn: {data.get('linkedin')}")
    if data.get("github"): contact_parts.append(f"GitHub: {data.get('github')}")
    
    run_contact = p_contact.add_run("  |  ".join(contact_parts))
    run_contact.font.size = Pt(9.5)
    run_contact.font.italic = True
    run_contact.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    p_contact.paragraph_format.space_after = Pt(14)
    
    # Helper to add section headers
    def add_section_header(title):
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(12)
        p_hdr.paragraph_format.space_after = Pt(4)
        p_hdr.paragraph_format.keep_with_next = True
        
        run_hdr = p_hdr.add_run(title.upper())
        run_hdr.font.size = Pt(12)
        run_hdr.font.bold = True
        run_hdr.font.color.rgb = RGBColor(0x1D, 0x4E, 0x89)
        
        # Add a subtle line below the section header
        p_hdr_border = doc.add_paragraph()
        p_hdr_border.paragraph_format.space_before = Pt(0)
        p_hdr_border.paragraph_format.space_after = Pt(6)
        # Using a sequence of underscores as a fallback separator
        run_border = p_hdr_border.add_run("_" * 60)
        run_border.font.size = Pt(6)
        run_border.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
        
    # 1. Summary
    if data.get("summary"):
        add_section_header("Professional Summary")
        p_sum = doc.add_paragraph(data.get("summary"))
        p_sum.paragraph_format.space_after = Pt(8)
        
    # 2. Preferred Job Role & Skills
    if data.get("skills") or data.get("preferred_role"):
        add_section_header("Core Competencies & Skills")
        if data.get("preferred_role"):
            p_role = doc.add_paragraph()
            run_role_lbl = p_role.add_run("Target Job Role: ")
            run_role_lbl.bold = True
            p_role.add_run(data.get("preferred_role"))
            p_role.paragraph_format.space_after = Pt(4)
            
        p_skills = doc.add_paragraph(data.get("skills"))
        p_skills.paragraph_format.space_after = Pt(8)
        
    # 3. Experience
    if data.get("experience"):
        add_section_header("Professional Experience")
        # Split experience by lines to represent paragraphs or bullet points
        exp_lines = data.get("experience", "").split("\n")
        for line in exp_lines:
            if line.strip():
                # Check if it looks like a job title or header (starts with bold markers or similar)
                if line.strip().startswith("-") or line.strip().startswith("*"):
                    bullet_text = line.strip()[1:].strip()
                    p_exp = doc.add_paragraph(style='List Bullet')
                    p_exp.add_run(bullet_text)
                else:
                    p_exp = doc.add_paragraph(line.strip())
                p_exp.paragraph_format.space_after = Pt(4)
                
    # 4. Education
    if data.get("education"):
        add_section_header("Education")
        edu_lines = data.get("education", "").split("\n")
        for line in edu_lines:
            if line.strip():
                if line.strip().startswith("-") or line.strip().startswith("*"):
                    bullet_text = line.strip()[1:].strip()
                    p_edu = doc.add_paragraph(style='List Bullet')
                    p_edu.add_run(bullet_text)
                else:
                    p_edu = doc.add_paragraph(line.strip())
                p_edu.paragraph_format.space_after = Pt(4)

    # 5. Projects
    if data.get("projects"):
        add_section_header("Projects")
        proj_lines = data.get("projects", "").split("\n")
        for line in proj_lines:
            if line.strip():
                if line.strip().startswith("-") or line.strip().startswith("*"):
                    bullet_text = line.strip()[1:].strip()
                    p_proj = doc.add_paragraph(style='List Bullet')
                    p_proj.add_run(bullet_text)
                else:
                    p_proj = doc.add_paragraph(line.strip())
                p_proj.paragraph_format.space_after = Pt(4)

    # 6. Certifications & Achievements
    if data.get("certifications") or data.get("achievements"):
        add_section_header("Certifications & Achievements")
        if data.get("certifications"):
            p_lbl = doc.add_paragraph()
            p_lbl.add_run("Certifications:").bold = True
            p_lbl.paragraph_format.space_after = Pt(2)
            cert_lines = data.get("certifications", "").split("\n")
            for line in cert_lines:
                if line.strip():
                    p_cert = doc.add_paragraph(style='List Bullet')
                    p_cert.add_run(line.strip()[1:].strip() if line.strip().startswith(("-", "*")) else line.strip())
                    p_cert.paragraph_format.space_after = Pt(3)
                    
        if data.get("achievements"):
            p_lbl = doc.add_paragraph()
            p_lbl.add_run("Key Achievements:").bold = True
            p_lbl.paragraph_format.space_before = Pt(6)
            p_lbl.paragraph_format.space_after = Pt(2)
            ach_lines = data.get("achievements", "").split("\n")
            for line in ach_lines:
                if line.strip():
                    p_ach = doc.add_paragraph(style='List Bullet')
                    p_ach.add_run(line.strip()[1:].strip() if line.strip().startswith(("-", "*")) else line.strip())
                    p_ach.paragraph_format.space_after = Pt(3)

    # Save to a memory stream
    output_stream = BytesIO()
    doc.save(output_stream)
    return output_stream.getvalue()
